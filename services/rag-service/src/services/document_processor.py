"""
Enhanced Document Processing Service for ENIAD RAG System
Supports PDF, Word, Images (OCR), and JSON files
"""

import os
import hashlib
import json
from typing import List, Dict, Any, Optional
from datetime import datetime
import logging

# Document processing libraries
import PyPDF2
import fitz  # PyMuPDF for better PDF handling
from PIL import Image
import pytesseract
from docx import Document as DocxDocument

# Database and models
from motor.motor_asyncio import AsyncIOMotorDatabase
from helpers.config import get_settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, db: AsyncIOMotorDatabase):
        self.db = db
        self.settings = get_settings()
        
        # Configure Tesseract for OCR
        if hasattr(self.settings, 'TESSERACT_CMD'):
            pytesseract.pytesseract.tesseract_cmd = self.settings.TESSERACT_CMD

    async def process_file(self, file_path: str, file_type: str, category: str = "general") -> Dict[str, Any]:
        """Process uploaded file and extract content"""
        try:
            # Generate file hash for deduplication
            file_hash = self._generate_file_hash(file_path)
            
            # Check if file already processed
            existing_doc = await self.db.documents.find_one({"file_hash": file_hash})
            if existing_doc:
                logger.info(f"File already processed: {file_path}")
                return existing_doc

            # Extract content based on file type
            content = ""
            metadata = {}
            
            if file_type == "application/pdf":
                content, metadata = await self._process_pdf(file_path)
            elif file_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                content, metadata = await self._process_word(file_path)
            elif file_type in ["image/jpeg", "image/png", "image/tiff"]:
                content, metadata = await self._process_image(file_path)
            elif file_type == "application/json":
                content, metadata = await self._process_json(file_path)
            elif file_type == "text/plain":
                content, metadata = await self._process_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            # Create document record
            document = {
                "title": os.path.basename(file_path),
                "content": content,
                "category": category,
                "language": self._detect_language(content),
                "file_type": file_type,
                "file_hash": file_hash,
                "file_path": file_path,
                "created_at": datetime.utcnow(),
                "metadata": metadata,
                "processed": True,
                "chunk_count": 0
            }

            # Insert document
            result = await self.db.documents.insert_one(document)
            document["_id"] = result.inserted_id

            logger.info(f"Successfully processed file: {file_path}")
            return document

        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise

    async def _process_pdf(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from PDF using PyMuPDF for better results"""
        content = ""
        metadata = {"pages": 0, "images": 0}
        
        try:
            # Use PyMuPDF for better text extraction
            doc = fitz.open(file_path)
            metadata["pages"] = len(doc)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                content += f"\n--- Page {page_num + 1} ---\n{text}"
                
                # Extract images if any
                image_list = page.get_images()
                metadata["images"] += len(image_list)
            
            doc.close()
            
        except Exception as e:
            logger.warning(f"PyMuPDF failed, falling back to PyPDF2: {str(e)}")
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    content += f"\n--- Page {page_num + 1} ---\n{text}"
        
        return content.strip(), metadata

    async def _process_word(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from Word documents"""
        try:
            doc = DocxDocument(file_path)
            content = ""
            metadata = {"paragraphs": 0, "tables": 0}
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
                    metadata["paragraphs"] += 1
            
            # Extract tables
            for table in doc.tables:
                metadata["tables"] += 1
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing Word document: {str(e)}")
            raise

    async def _process_image(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_path)
            metadata = {
                "width": image.width,
                "height": image.height,
                "format": image.format
            }
            
            # Perform OCR
            ocr_config = '--oem 3 --psm 6'
            content = pytesseract.image_to_string(
                image, 
                lang=self.settings.OCR_LANGUAGES,
                config=ocr_config
            )
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            raise

    async def _process_json(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process JSON files (like question datasets)"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            content = ""
            metadata = {"entries": 0, "structure": type(data).__name__}
            
            if isinstance(data, list):
                metadata["entries"] = len(data)
                for item in data:
                    if isinstance(item, dict):
                        # Extract question-answer pairs
                        if "question" in item and "answer" in item:
                            content += f"Q: {item['question']}\nA: {item['answer']}\n\n"
                        else:
                            content += json.dumps(item, ensure_ascii=False) + "\n"
                    else:
                        content += str(item) + "\n"
            elif isinstance(data, dict):
                content = json.dumps(data, ensure_ascii=False, indent=2)
                metadata["entries"] = len(data)
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing JSON file: {str(e)}")
            raise

    async def _process_text(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            metadata = {
                "lines": len(content.split('\n')),
                "characters": len(content),
                "words": len(content.split())
            }
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing text file: {str(e)}")
            raise

    def _generate_file_hash(self, file_path: str) -> str:
        """Generate SHA-256 hash of file for deduplication"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def _detect_language(self, content: str) -> str:
        """Simple language detection based on content"""
        # Arabic detection
        arabic_chars = sum(1 for char in content if '\u0600' <= char <= '\u06FF')
        if arabic_chars > len(content) * 0.1:  # More than 10% Arabic characters
            return "ar"
        
        # French detection (basic)
        french_words = ['le', 'la', 'les', 'de', 'du', 'des', 'et', 'est', 'une', 'un']
        content_lower = content.lower()
        french_count = sum(1 for word in french_words if word in content_lower)
        
        if french_count > 3:
            return "fr"
        
        return self.settings.DEFAULT_LANG

    async def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        try:
            total_docs = await self.db.documents.count_documents({})
            
            # Group by category
            pipeline = [
                {"$group": {"_id": "$category", "count": {"$sum": 1}}},
                {"$sort": {"count": -1}}
            ]
            categories = await self.db.documents.aggregate(pipeline).to_list(None)
            
            # Group by language
            pipeline = [
                {"$group": {"_id": "$language", "count": {"$sum": 1}}},
                {"$sort": {"count": -1}}
            ]
            languages = await self.db.documents.aggregate(pipeline).to_list(None)
            
            return {
                "total_documents": total_docs,
                "categories": categories,
                "languages": languages,
                "last_updated": datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error getting document stats: {str(e)}")
            return {"error": str(e)}
